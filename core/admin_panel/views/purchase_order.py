import json

from django.contrib.auth.mixins import LoginRequiredMixin
from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse, HttpResponse
from django.contrib import messages
from django.db.models import Q
from django.urls import reverse
from django.views.decorators.http import require_http_methods
from django.views import View
from django.utils import timezone

from core.admin_panel.models.purchase_order import (
    PurchaseOrder,
    PurchaseOrderOperation,
    PurchaseOrderAccessory,
    PurchaseOrderStatus,
    AccessoryType
)
from core.operations_panel.models import Operation, Client, Driver, Supplier
from apps.telegram_bots.models import TelegramUser
from core.system.views import AdminListView


class PurchaseOrderListView(AdminListView):
    """Vista principal del listado de órdenes de compra"""
    template_name = 'admin_panel/purchase_order_management.html'
    model = PurchaseOrder

    def post(self, request):
        if request.POST.get('action') == 'searchdata':
            return self._search_data(request)
        return JsonResponse({'error': 'Acción no válida'})

    def _search_data(self, request):
        """Maneja la búsqueda para DataTables"""
        search = request.POST.get('search', '')
        start = int(request.POST.get('start', 0))
        length = int(request.POST.get('length', 10))

        queryset = PurchaseOrder.objects.select_related('supplier', 'driver').all()

        if search:
            queryset = queryset.filter(
                Q(folio__icontains=search) |
                Q(supplier__business_name__icontains=search) |
                Q(driver__name__icontains=search)
            )

        total = queryset.count()
        orders = queryset[start:start + length]

        data = []
        for order in orders:
            order.calculate_totals()
            data.append({
                'id': order.id,
                'folio': order.folio,
                'client': order.supplier.business_name if order.supplier else 'N/A',
                'driver': order.driver.name if order.driver else 'N/A',
                'status': order.get_status_display(),
                'total': f'${order.total:,.2f}',
                'created_at': order.created_at.strftime('%d/%m/%Y'),
            })

        return JsonResponse({
            'recordsTotal': PurchaseOrder.objects.count(),
            'recordsFiltered': total,
            'data': data
        })

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['title'] = 'Órdenes de Compra'
        context['datatable_keys'] = [
                'folio', 'clients', 'driver', 'status',
                'total', 'created_at'
            ]
        context['datatable_actions'] = True
        return context

def purchase_order_create(request):
    """Vista para crear nueva orden de compra"""

    if request.method == 'GET':
        clients = Supplier.objects.all()
        drivers = Driver.objects.all()

        context = {
            'title': 'Nueva Orden de Compra',
            'clients': clients,
            'drivers': drivers,
            'accessory_types': AccessoryType.choices,
        }
        return render(request, 'admin_panel/purchase_order_form.html', context)

    elif request.method == 'POST':
        try:
            data = json.loads(request.body)

            order = PurchaseOrder.objects.create(
                supplier_id=data['client_id'],
                driver_id=data.get('driver_id') or None,
                notes=data.get('notes', ''),
            )

            # Agregar operaciones
            for op_data in data.get('operations', []):
                op_id = op_data.get('id') if isinstance(op_data, dict) else op_data

                operation = get_object_or_404(Operation, id=op_id)

                # Si quieres actualizar el costo base de la operación
                if isinstance(op_data, dict) and 'cost' in op_data:
                    operation.total = op_data.get('cost') or 0
                    operation.save(update_fields=['total'])

                PurchaseOrderOperation.objects.create(
                    purchase_order=order,
                    operation=operation
                )

            # Agregar accesorios
            for acc_data in data.get('accessories', []):
                operation_id = acc_data.get('operation_id')

                operation = get_object_or_404(Operation, id=operation_id)

                PurchaseOrderAccessory.objects.create(
                    purchase_order=order,
                    operation=operation,
                    type=acc_data.get('type'),
                    description=acc_data.get('description', ''),
                    quantity=acc_data.get('quantity') or 0,
                    unit_price=acc_data.get('unit_price') or 0
                )

            # Impuestos manuales
            if hasattr(order, 'tax'):
                order.tax = data.get('tax') or 0
                order.save(update_fields=['tax'])

            order.calculate_totals()

            return JsonResponse({
                'success': True,
                'redirect': reverse('admin_panel:purchase_order_detail', args=[order.id])
            })

        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            })


def get_operations_by_filter(request):
    """API para obtener operaciones filtradas por cliente y opcionalmente por folio"""
    supplier_id = request.GET.get('client_id')
    folio_search = request.GET.get('folio_search', '').strip()

    if not supplier_id:
        return JsonResponse({'operations': []})

    queryset = Operation.objects.filter(supplier_id=supplier_id).order_by('-folio')

    # Filtrar por folio si se proporciona
    if folio_search:
        queryset = queryset.filter(
            Q(folio__icontains=folio_search) if hasattr(Operation, 'folio')
            else Q(id__icontains=folio_search)
        )

    operations = []
    for op in queryset[:10]:  # Limitar a 50 resultados
        print(op.total)
        operations.append({
            'id': str(op.id),
            'origin': str(op.route.initial_location) if op.route and op.route.initial_location else 'N/A',
            'destination': str(op.route.destination_location) if op.route and op.route.destination_location else 'N/A',
            'cost': float(getattr(op, 'total', 0) or 0),
            'date': op.cargo_appointment.strftime('%d/%m/%Y') if op.cargo_appointment else 'N/A',
            'folio': getattr(op, 'folio', f'OP-{op.id}')
        })

    return JsonResponse({'operations': operations})


def purchase_order_detail(request, order_id):
    """Vista de detalle de orden de compra"""
    order = get_object_or_404(PurchaseOrder, id=order_id)

    context = {
        'title': f'Orden de Compra {order.folio}',
        'order': order,
        'operations': order.operations.select_related('operation').all(),
        'accessories': order.accessories.select_related('operation').all(),
        'status_choices': PurchaseOrderStatus.choices,
    }
    return render(request, 'admin_panel/purchase_order_details.html', context)


@require_http_methods(["POST"])
def purchase_order_update_status(request, order_id):
    """Actualizar el estado de una orden de compra"""
    order = get_object_or_404(PurchaseOrder, id=order_id)
    new_status = request.POST.get('status')

    if new_status in dict(PurchaseOrderStatus.choices):
        old_status = order.status
        order.status = new_status

        # Actualizar fechas según el estado
        if new_status == PurchaseOrderStatus.ENVIADA:
            order.sent_date = timezone.now()
        elif new_status == PurchaseOrderStatus.EN_ESPERA:
            order.waiting_date = timezone.now()
        elif new_status == PurchaseOrderStatus.APROBADA:
            order.approved_date = timezone.now()
        elif new_status == PurchaseOrderStatus.PAGADA:
            order.paid_date = timezone.now()

        order.save()

        messages.success(
            request,
            f'Estado actualizado de {old_status} a {order.get_status_display()}'
        )
    else:
        messages.error(request, 'Estado no válido')

    return redirect('admin_panel:purchase_order_detail', order_id)


def purchase_order_generate_pdf(request, order_id):
    """Generar PDF de la orden de compra"""
    order = get_object_or_404(PurchaseOrder, id=order_id)

    # Por ahora, respuesta simple - puedes implementar generación de PDF más tarde
    from django.template.loader import render_to_string

    html_content = render_to_string(
        'admin_panel/purchase_order/pdf_template.html',
        {'order': order}
    )

    response = HttpResponse(html_content, content_type='text/html')
    response['Content-Disposition'] = f'inline; filename="{order.folio}.html"'

    return response



import os
from decimal import Decimal
from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    borders = OxmlElement('w:tblBorders')

    for border_name in [
        'top',
        'left',
        'bottom',
        'right',
        'insideH',
        'insideV'
    ]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)

    tblPr.append(borders)

def money(value):
    return f"{Decimal(value or 0):,.2f}"


def replace_all(document, replacements):
    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))


def insert_after(anchor_xml, block):
    element = block._element
    anchor_xml.addnext(element)
    return element

def apply_table_style(table):
    try:
        table.style = "Table Grid"
    except KeyError:
        try:
            table.style = "Tabla con cuadrícula"
        except KeyError:
            pass


def purchase_order_generate_docx(request, order_id):
    order = get_object_or_404(PurchaseOrder, id=order_id)

    template_path = os.path.join(
        settings.BASE_DIR,
        "static",
        "templates",
        "PlantillaOC.docx"
    )

    document = Document(template_path)

    operations_rel = PurchaseOrderOperation.objects.filter(
        purchase_order=order
    ).select_related("operation", "operation__route")

    accessories = PurchaseOrderAccessory.objects.filter(
        purchase_order=order
    ).select_related("operation")

    subtotal_operations = Decimal("0")
    subtotal_accessories = Decimal("0")

    for rel in operations_rel:
        subtotal_operations += Decimal(getattr(rel.operation, "total", 0) or 0)

    for acc in accessories:
        subtotal_accessories += Decimal(acc.subtotal or 0)

    tax = Decimal(getattr(order, "tax", 0) or 0)
    total = subtotal_operations + subtotal_accessories + tax

    replacements = {
        "<FOLIO>": order.folio,
        "<FECHA_EXP>": order.created_at.strftime("%d/%m/%Y"),
        "<NOMBRE_CLIENTE>": str(order.supplier),
        "<CLIENTE_ID>": str(order.supplier_id),
        "<DRIVER>": str(order.driver) if order.driver else "N/A",
        "<ESTATUS>": getattr(order, "status", "N/A"),
        "<NOTAS>": order.notes or "Sin notas",
        "<SUBTOTAL_OPERACIONES>": money(subtotal_operations),
        "<SUBTOTAL_ACCESORIOS>": money(subtotal_accessories),
        "<IMPUESTOS>": money(tax),
        "<TOTAL>": money(total),
    }

    replace_all(document, replacements)

    marker_paragraph = None
    for paragraph in document.paragraphs:
        if "<DETALLE_OPERACIONES>" in paragraph.text:
            marker_paragraph = paragraph
            paragraph.text = ""
            break

    anchor = marker_paragraph._p

    for index, rel in enumerate(operations_rel, start=1):
        op = rel.operation
        route = getattr(op, "route", None)

        op_accessories = accessories.filter(operation=op)
        op_base_total = Decimal(getattr(op, "total", 0) or 0)
        op_accessories_total = sum(
            Decimal(acc.subtotal or 0)
            for acc in op_accessories
        )
        op_total = op_base_total + op_accessories_total

        title = document.add_paragraph()
        title_run = title.add_run(
            f"Folio: {getattr(op, 'folio', f'OP-{op.id}')}"
        )
        title_run.bold = True
        title_run.font.size = Pt(11)
        anchor = insert_after(anchor, title)

        info_table = document.add_table(rows=4, cols=2)
        set_table_borders(info_table)

        info_rows = [
            ("Origen", str(route.initial_location) if route and route.initial_location else "N/A"),
            ("Destino", str(route.destination_location) if route and route.destination_location else "N/A"),
            ("Fecha", op.cargo_appointment.strftime("%d/%m/%Y") if getattr(op, "cargo_appointment", None) else "N/A"),
            ("Costo base", f"${money(op_base_total)}"),
        ]

        for row, (label, value) in zip(info_table.rows, info_rows):
            row.cells[0].text = label
            row.cells[1].text = value

        anchor = insert_after(anchor, info_table)

        acc_title = document.add_paragraph()
        acc_title.add_run("Accesorios").bold = True
        anchor = insert_after(anchor, acc_title)

        acc_table = document.add_table(rows=1, cols=5)
        set_table_borders(acc_table)

        headers = ["Tipo", "Descripcion", "Cantidad", "Costo", "Subtotal"]
        for cell, header in zip(acc_table.rows[0].cells, headers):
            cell.text = header

        if op_accessories.exists():
            for acc in op_accessories:
                row = acc_table.add_row().cells
                row[0].text = acc.type
                row[1].text = acc.description
                row[2].text = str(acc.quantity)
                row[3].text = f"${money(acc.unit_price)}"
                row[4].text = f"${money(acc.subtotal)}"
        else:
            row = acc_table.add_row().cells
            row[0].text = "Sin accesorios"
            row[1].text = ""
            row[2].text = ""
            row[3].text = ""
            row[4].text = "$0.00"

        anchor = insert_after(anchor, acc_table)

        total_table = document.add_table(rows=3, cols=2)
        set_table_borders(total_table)

        totals = [
            ("Subtotal operacion", f"${money(op_base_total)}"),
            ("Subtotal accesorios", f"${money(op_accessories_total)}"),
            ("Total operacion", f"${money(op_total)}"),
        ]

        for row, (label, value) in zip(total_table.rows, totals):
            row.cells[0].text = label
            row.cells[1].text = value

        anchor = insert_after(anchor, total_table)

        spacer = document.add_paragraph("")
        anchor = insert_after(anchor, spacer)

        for row in acc_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)

        for cell in acc_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    filename = f"orden-pago-{order.folio}.docx"
    output_path = os.path.join(settings.BASE_DIR, "static", filename)
    document.save(output_path)

    with open(output_path, "rb") as doc_file:
        response = HttpResponse(
            doc_file.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response